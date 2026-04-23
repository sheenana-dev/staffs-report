"""Generate weekly report Google Doc templates as .docx files."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


def style_doc(doc):
    """Apply consistent styling to the document."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(10)

    for i in range(1, 4):
        heading = doc.styles[f"Heading {i}"]
        heading.font.name = "Arial"
        heading.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)


def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + rows, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            run = paragraph.runs[0]
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = "Arial"

    # Empty data rows
    for row_idx in range(1, 1 + rows):
        for col_idx in range(len(headers)):
            cell = table.rows[row_idx].cells[col_idx]
            for paragraph in cell.paragraphs:
                paragraph.style.font.size = Pt(9)

    return table


def add_field(doc, label, default="___"):
    """Add a labeled field line."""
    p = doc.add_paragraph()
    run = p.add_run(f"{label} ")
    run.bold = True
    run.font.size = Pt(10)
    p.add_run(default).font.size = Pt(10)


def add_numbered_blanks(doc, count):
    """Add numbered blank lines."""
    for i in range(1, count + 1):
        doc.add_paragraph(f"{i}.", style="List Number")


def add_section_break(doc):
    doc.add_paragraph("—" * 40)


# ============================================================
# IT DEPARTMENT
# ============================================================
def create_it_report():
    doc = Document()
    style_doc(doc)

    doc.add_heading("IT Department — Weekly Report", level=1)
    doc.add_paragraph()
    add_field(doc, "Lead Name:")
    add_field(doc, "Week:", "[Date] to [Date]")
    add_section_break(doc)

    doc.add_heading("1. What We Shipped This Week", level=2)
    add_table(doc, ["Feature / Task", "Status", "Notes"], 4)
    doc.add_paragraph("Status options: Done / In Progress / Blocked", style="Intense Quote")

    doc.add_heading("2. Bugs & Issues", level=2)
    add_table(doc, ["Bug / Issue", "Severity", "Status", "Resolved Date"], 3)
    doc.add_paragraph("Severity options: Critical / High / Low", style="Intense Quote")
    doc.add_paragraph()
    add_field(doc, "Critical bugs resolved this week:")
    add_field(doc, "Open critical bugs:")

    doc.add_heading("3. Platform Health", level=2)
    add_field(doc, "App uptime this week:", "____%")
    add_field(doc, "API uptime this week:", "____%")
    add_field(doc, "Incidents / outages:", "None")
    add_field(doc, "Deployments this week:")

    doc.add_heading("4. UI/UX Updates", level=2)
    add_table(doc, ["Design Task", "Status", "Handed Off to Dev?"], 3)

    doc.add_heading("5. Sprint Progress", level=2)
    add_field(doc, "Current sprint:", "Sprint ___ ([dates])")
    add_field(doc, "Sprint completion so far:", "____%")
    add_field(doc, "At risk items:")

    doc.add_heading("6. App Store", level=2)
    add_field(doc, "iOS rating:")
    add_field(doc, "Android rating:")
    add_field(doc, "Notable reviews this week:")

    doc.add_heading("7. Cross-Department Needs", level=2)
    add_table(doc, ["I Need From...", "What", "Urgency"], 3)
    doc.add_paragraph("Urgency options: This week / Next week / Soon", style="Intense Quote")

    doc.add_heading("8. Blockers", level=2)
    doc.add_paragraph("What's slowing the team down right now?")
    add_numbered_blanks(doc, 2)

    doc.add_heading("9. Next Week Priorities", level=2)
    add_numbered_blanks(doc, 3)

    doc.add_heading("10. Team Notes", level=2)
    doc.add_paragraph("Anything about the team I should know (workload, morale, time off)?")
    doc.add_paragraph()

    add_section_break(doc)
    p = doc.add_paragraph()
    run = p.add_run("Submit every Friday by 4 PM in the shared Google Docs folder.")
    run.italic = True
    run.font.size = Pt(9)

    path = os.path.join(OUTPUT_DIR, "Weekly Report - IT Department.docx")
    doc.save(path)
    print(f"Created: {path}")


# ============================================================
# DESIGN DEPARTMENT
# ============================================================
def create_design_report():
    doc = Document()
    style_doc(doc)

    doc.add_heading("Design Department — Weekly Report", level=1)
    doc.add_paragraph()
    add_field(doc, "Lead Name:")
    add_field(doc, "Week:", "[Date] to [Date]")
    add_section_break(doc)

    doc.add_heading("1. Completed Requests This Week", level=2)
    add_table(doc, ["Request", "Requested By", "Type", "Delivered Date"], 4)
    doc.add_paragraph("Type options: Merchant / Campaign / Brand / Other", style="Intense Quote")
    doc.add_paragraph()
    add_field(doc, "Total requests completed this week:")

    doc.add_heading("2. In Progress", level=2)
    add_table(doc, ["Request", "Requested By", "Expected Delivery", "Notes"], 3)

    doc.add_heading("3. Backlog", level=2)
    add_table(doc, ["Request", "Requested By", "Priority", "Waiting Since"], 3)
    doc.add_paragraph("Priority options: High / Medium / Low", style="Intense Quote")
    doc.add_paragraph()
    add_field(doc, "Total backlog size:", "___ requests")

    doc.add_heading("4. Key Deliverables This Week", level=2)
    doc.add_paragraph("List the most important design assets delivered:")
    add_numbered_blanks(doc, 3)

    doc.add_heading("5. Merchant Materials", level=2)
    add_field(doc, "Merchant collateral produced this week:", "___ (posters, banners, menus, etc.)")
    add_field(doc, "Merchants served:")

    doc.add_heading("6. Campaign / Promotional Assets", level=2)
    add_field(doc, "Campaign visuals delivered:", "___ pieces")
    add_field(doc, "For which campaigns:")
    add_field(doc, "App store visuals updated:", "Yes / No")

    doc.add_heading("7. Cross-Department Needs", level=2)
    add_table(doc, ["I Need From...", "What", "Urgency"], 3)

    doc.add_heading("8. Blockers", level=2)
    doc.add_paragraph("What's slowing the team down right now?")
    add_numbered_blanks(doc, 2)

    doc.add_heading("9. Next Week Priorities", level=2)
    add_numbered_blanks(doc, 3)

    doc.add_heading("10. Team Notes", level=2)
    doc.add_paragraph("Anything about the team I should know (workload, morale, time off)?")
    doc.add_paragraph()

    add_section_break(doc)
    p = doc.add_paragraph()
    run = p.add_run("Submit every Friday by 4 PM in the shared Google Docs folder.")
    run.italic = True
    run.font.size = Pt(9)

    path = os.path.join(OUTPUT_DIR, "Weekly Report - Design Department.docx")
    doc.save(path)
    print(f"Created: {path}")


# ============================================================
# SALES DEPARTMENT
# ============================================================
def create_sales_report():
    doc = Document()
    style_doc(doc)

    doc.add_heading("Sales Department — Weekly Report", level=1)
    doc.add_paragraph()
    add_field(doc, "Lead Name:")
    add_field(doc, "Week:", "[Date] to [Date]")
    add_section_break(doc)

    # --- Sales Team ---
    doc.add_heading("1. Sales Team (Merchant Acquisition)", level=2)

    doc.add_heading("New Merchants This Week", level=3)
    add_table(doc, ["Merchant Name", "Location", "Category", "Status"], 4)
    doc.add_paragraph("Status options: Pitched / Negotiating / Signed", style="Intense Quote")
    doc.add_paragraph()
    add_field(doc, "Merchants contacted / pitched this week:")
    add_field(doc, "Merchants signed this week:")
    add_field(doc, "Total merchants in pipeline:")

    doc.add_heading("Banner Ad Sales", level=3)
    add_table(doc, ["Client / Merchant", "Banner Type", "Revenue (PHP)", "Status"], 3)
    doc.add_paragraph("Status options: Pitched / Sold / Live", style="Intense Quote")
    doc.add_paragraph()
    add_field(doc, "Banner revenue closed this week:", "PHP ___")

    # --- Onboarding Ops ---
    doc.add_heading("2. Onboarding Ops", level=2)
    add_table(doc, ["Merchant Name", "Date Signed", "Onboarding Status", "HubSpot Updated?", "Added to App?"], 3)
    doc.add_paragraph("Status options: Collecting Info / In Progress / Complete", style="Intense Quote")
    doc.add_paragraph()
    add_field(doc, "Merchants fully onboarded this week:")
    add_field(doc, "Merchants pending (waiting for merchant info):")
    add_field(doc, "Average onboarding time this week:", "___ days")

    doc.add_heading("Onboarding Checklist (per merchant)", level=3)
    checklist_items = [
        "Photos collected",
        "Operating hours confirmed",
        "Contract signed",
        "HubSpot record created",
        "Merchant added to B-Ticket app",
        "Coupons set up in app",
        "Merchant notified they're live",
    ]
    for item in checklist_items:
        doc.add_paragraph(f"[ ]  {item}")

    # --- Account Management Ops ---
    doc.add_heading("3. Account Management Ops", level=2)

    doc.add_heading("Coupon Performance Review", level=3)
    add_table(doc, ["Merchant Name", "Coupons Active", "Redemptions This Month", "Action Needed?"], 3)
    doc.add_paragraph("Action options: None / Swap coupon / Follow up", style="Intense Quote")
    doc.add_paragraph()
    add_field(doc, "Merchants reviewed this week:")
    add_field(doc, "Coupon swaps requested:")
    add_field(doc, "Merchants at risk (low redemptions):")

    # --- Pipeline Summary ---
    doc.add_heading("4. Pipeline Summary", level=2)
    pipeline_table = add_table(doc, ["Stage", "Count"], 6)
    stages = ["New leads", "Pitched / In talks", "Negotiating", "Contract sent", "Signed (pending onboarding)", "Live on app"]
    for i, stage in enumerate(stages):
        pipeline_table.rows[i + 1].cells[0].text = stage

    # --- Cross-Dept ---
    doc.add_heading("5. Cross-Department Needs", level=2)
    add_table(doc, ["I Need From...", "What", "Urgency"], 3)

    doc.add_heading("6. Blockers", level=2)
    doc.add_paragraph("What's slowing the team down right now?")
    add_numbered_blanks(doc, 2)

    doc.add_heading("7. Next Week Priorities", level=2)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Sales:").bold = True
    add_numbered_blanks(doc, 2)
    p = doc.add_paragraph()
    p.add_run("Onboarding Ops:").bold = True
    add_numbered_blanks(doc, 2)
    p = doc.add_paragraph()
    p.add_run("Account Management Ops:").bold = True
    add_numbered_blanks(doc, 2)

    doc.add_heading("8. Team Notes", level=2)
    doc.add_paragraph("Anything about the team I should know (workload, morale, time off)?")
    doc.add_paragraph()

    add_section_break(doc)
    p = doc.add_paragraph()
    run = p.add_run("Submit every Friday by 4 PM in the shared Google Docs folder.")
    run.italic = True
    run.font.size = Pt(9)

    path = os.path.join(OUTPUT_DIR, "Weekly Report - Sales Department.docx")
    doc.save(path)
    print(f"Created: {path}")


# ============================================================
# MARKETING DEPARTMENT
# ============================================================
def create_marketing_report():
    doc = Document()
    style_doc(doc)

    doc.add_heading("Marketing Department — Weekly Report", level=1)
    doc.add_paragraph()
    add_field(doc, "Lead Name:")
    add_field(doc, "Week:", "[Date] to [Date]")
    add_section_break(doc)

    doc.add_heading("1. Key Numbers This Week", level=2)
    metrics_table = add_table(doc, ["Metric", "This Week", "Last Week", "Change (+/-)"], 7)
    metrics = ["App downloads (total)", "iOS downloads", "Android downloads", "Huawei downloads", "Active users", "New subscribers", "Coupon redemptions"]
    for i, metric in enumerate(metrics):
        metrics_table.rows[i + 1].cells[0].text = metric

    doc.add_heading("2. Content Published This Week", level=2)
    add_table(doc, ["Content", "Platform", "Date Posted", "Performance"], 4)
    doc.add_paragraph()
    add_field(doc, "Total content pieces published:")

    doc.add_heading("3. Social Media Overview", level=2)
    social_table = add_table(doc, ["Platform", "Followers", "Posts This Week", "Top Performing Post"], 4)
    platforms = ["Facebook", "Instagram", "TikTok", "Twitter/X"]
    for i, platform in enumerate(platforms):
        social_table.rows[i + 1].cells[0].text = platform

    doc.add_heading("4. Campaigns", level=2)

    doc.add_heading("Active Campaigns", level=3)
    add_table(doc, ["Campaign Name", "Goal", "Start - End Date", "Status", "Results So Far"], 3)
    doc.add_paragraph("Status options: Running / Paused / Ended", style="Intense Quote")

    doc.add_paragraph()
    add_field(doc, "Campaigns launched this week:")
    add_field(doc, "Campaigns ending soon:")

    doc.add_heading("5. Paid Advertising", level=2)
    add_table(doc, ["Platform", "Budget (PHP)", "Spent This Week", "Results"], 2)

    doc.add_heading("6. Merchant Promotion Support", level=2)
    doc.add_paragraph("How are we helping Sales drive redemptions for merchants this week?")
    add_table(doc, ["Merchant / Campaign", "What We Did", "Platform", "Results"], 3)

    doc.add_heading("7. Cross-Department Needs", level=2)
    add_table(doc, ["I Need From...", "What", "Urgency"], 3)

    doc.add_heading("8. Blockers", level=2)
    doc.add_paragraph("What's slowing the team down right now?")
    add_numbered_blanks(doc, 2)

    doc.add_heading("9. Next Week Priorities", level=2)
    add_numbered_blanks(doc, 3)

    doc.add_heading("10. Team Notes", level=2)
    doc.add_paragraph("Anything about the team I should know (workload, morale, time off)?")
    doc.add_paragraph()

    add_section_break(doc)
    p = doc.add_paragraph()
    run = p.add_run("Submit every Friday by 4 PM in the shared Google Docs folder.")
    run.italic = True
    run.font.size = Pt(9)

    path = os.path.join(OUTPUT_DIR, "Weekly Report - Marketing Department.docx")
    doc.save(path)
    print(f"Created: {path}")


if __name__ == "__main__":
    create_it_report()
    create_design_report()
    create_sales_report()
    create_marketing_report()
    print("\nAll 4 weekly report templates created!")
